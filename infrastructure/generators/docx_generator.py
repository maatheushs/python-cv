from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from application.interfaces.resume_generator import ResumeGenerator
from domain.entities.resume import Resume

_HEADERS = {
    "en": {
        "summary": "Professional Summary",
        "experience": "Professional Experience",
        "skills": "Technical Skills",
        "education": "Education",
        "certifications": "Certifications",
    },
    "pt": {
        "summary": "Resumo Profissional",
        "experience": "Experiência Profissional",
        "skills": "Habilidades Técnicas",
        "education": "Formação Acadêmica",
        "certifications": "Certificações",
    },
}


class DocxResumeGenerator(ResumeGenerator):
    def _setup_styles(self):
        style = self._doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        for section in self._doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

    def generate(self, resume: Resume, output_path: str) -> None:
        self._doc = Document()
        self._setup_styles()
        self._h = _HEADERS.get(resume.locale, _HEADERS["en"])
        self._add_header(resume)
        self._add_summary(resume.summary)
        self._add_experience(resume)
        self._add_skills(resume)
        self._add_education(resume)
        if resume.certifications:
            self._add_certifications(resume)
        self._doc.save(output_path)

    def get_text_content(self) -> str:
        return "\n".join(p.text for p in self._doc.paragraphs if p.text.strip())

    def _section_title(self, title: str):
        run = self._doc.add_paragraph().add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        self._doc.add_paragraph().add_run("_" * 80).font.size = Pt(8)

    def _add_header(self, resume: Resume):
        p = resume.personal

        name_p = self._doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(p.name)
        name_run.bold = True
        name_run.font.size = Pt(16)

        title_p = self._doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.add_run(p.title).font.size = Pt(12)

        contact = f"{p.phone} | {p.email} | {p.location}"
        if p.linkedin:
            contact += f" | {p.linkedin}"
        contact_p = self._doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run(contact).font.size = Pt(10)

        self._doc.add_paragraph()

    def _add_summary(self, summary: str):
        self._section_title(self._h["summary"])
        self._doc.add_paragraph(summary).paragraph_format.space_after = Pt(10)
        self._doc.add_paragraph()

    def _add_experience(self, resume: Resume):
        self._section_title(self._h["experience"])
        for exp in resume.experience:
            pos_run = self._doc.add_paragraph().add_run(exp.position)
            pos_run.bold = True
            pos_run.font.size = Pt(11)

            self._doc.add_paragraph(
                f"{exp.company} | {exp.location} | {exp.start_date} - {exp.end_date}"
            ).add_run("").font.size = Pt(10)

            for item in exp.description:
                bullet = self._doc.add_paragraph(item, style="List Bullet")
                bullet.paragraph_format.left_indent = Inches(0.25)
                bullet.paragraph_format.space_after = Pt(2)

            self._doc.add_paragraph()

    def _add_skills(self, resume: Resume):
        self._section_title(self._h["skills"])
        for category, skill_list in resume.skills.items():
            if not skill_list:
                continue
            p = self._doc.add_paragraph()
            p.add_run(f"{category.replace('_', ' ').title()}: ").bold = True
            p.add_run(", ".join(skill_list) if isinstance(skill_list, list) else skill_list)
            p.paragraph_format.space_after = Pt(4)
        self._doc.add_paragraph()

    def _add_education(self, resume: Resume):
        self._section_title(self._h["education"])
        for edu in resume.education:
            p = self._doc.add_paragraph()
            p.add_run(edu.degree).bold = True
            self._doc.add_paragraph(
                f"{edu.institution} | {edu.location} | {edu.start_date} - {edu.end_date}"
            ).paragraph_format.space_after = Pt(6)
        self._doc.add_paragraph()

    def _add_certifications(self, resume: Resume):
        self._section_title(self._h["certifications"])
        for cert in resume.certifications:
            p = self._doc.add_paragraph()
            p.add_run(cert["name"]).bold = True
            details = f" - {cert['issuer']}"
            if "level" in cert:
                details += f" | {cert['level']}"
            if "date" in cert:
                details += f" | {cert['date']}"
            p.add_run(details)
            p.paragraph_format.space_after = Pt(4)
